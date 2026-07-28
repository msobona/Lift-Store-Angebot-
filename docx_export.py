"""Word-Export: Software-Anhang mit Platzhaltern befüllen (ein Dokument)."""

from __future__ import annotations

import io
import json
import re
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional
from xml.sax.saxutils import escape as xml_escape

from docxtpl import DocxTemplate

BASE_DIR = Path(__file__).resolve().parent
FIELD_MAP_FILE = BASE_DIR / "data" / "docx_field_map.json"
SSI_CONTACTS_FILE = BASE_DIR / "data" / "ssi_contacts.json"

# Word-Inhaltssteuerelemente (ComboBox / Text) in Angebot Logimat DE.docx
_SDT_TEXT_MAP = {
    "Name (Innendienst)": "ssi_kontakt_1_name",
    "Text Innendienst": "ssi_kontakt_1_details",
    "Name (Aussendienst)": "ssi_kontakt_2_name",
    "Text Aussendienst": "ssi_kontakt_2_details",
    "Unterschrift links": "signatur_1_name",
    "Funktion links": "signatur_1_details",
    "Unterschrift rechts": "signatur_2_name",
    "Funktion rechts": "signatur_2_details",
}

# Primär: Logimat-Vorlage (Titelseite mit Bild + Platzhalter).
# Fallbacks: vollständige Anhang-Vorlage / schlanke Platzhalter-Vorlage.
PLACEHOLDER_CANDIDATES = [
    BASE_DIR / "docs" / "templates" / "Angebot Logimat DE.docx",
    BASE_DIR / "docs" / "templates" / "anhang_angebot_vorlage.docx",
    BASE_DIR / "docs" / "templates" / "angebot_platzhalter.docx",
]

INSTANCE_DISPLAY_NAMES = {
    "basic": "Basis-Instanz",
    "advanced": "Advanced-Instanz",
    "Basic Instance": "Basis-Instanz",
    "Advanced Instance": "Advanced-Instanz",
}


def find_placeholder_template() -> Path:
    for path in PLACEHOLDER_CANDIDATES:
        if path.exists():
            return path
    raise FileNotFoundError(
        "Word-Vorlage nicht gefunden. Erwartet z.B.:\n"
        "  docs/templates/anhang_angebot_vorlage.docx"
    )


def ensure_annex_template() -> Path:
    """Liefert die aktive Word-Vorlage (Anhang / Fallback / Logimat)."""
    return find_placeholder_template()


def load_field_map() -> Dict[str, Any]:
    if FIELD_MAP_FILE.exists():
        return json.loads(FIELD_MAP_FILE.read_text(encoding="utf-8"))
    return {}


def _load_ssi_contacts_catalog() -> List[Dict[str, Any]]:
    if not SSI_CONTACTS_FILE.exists():
        return []
    try:
        data = json.loads(SSI_CONTACTS_FILE.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return []
    return list(data.get("contacts") or [])


def _lookup_ssi_contact(
    *, contact_id: str = "", name: str = ""
) -> Dict[str, Any]:
    cid = (contact_id or "").strip()
    needle = (name or "").strip().lower()
    for row in _load_ssi_contacts_catalog():
        if cid and row.get("id") == cid:
            return dict(row)
    if needle:
        for row in _load_ssi_contacts_catalog():
            if str(row.get("name") or "").strip().lower() == needle:
                return dict(row)
    return {}


def _money(value: Any, currency: str = "CHF") -> str:
    try:
        amount = float(value or 0)
    except (TypeError, ValueError):
        amount = 0.0
    # Immer 2 Dezimalstellen (Verkaufspreise in der Kalkulation auf ganze CHF aufgerundet)
    formatted = f"{amount:,.2f}".replace(",", "X").replace(".", ",").replace("X", "'")
    return f"{currency} {formatted}"


def _fmt_num(value: Any) -> str:
    if value is None or value == "":
        return ""
    try:
        num = float(value)
    except (TypeError, ValueError):
        return str(value)
    if abs(num - int(num)) < 1e-9:
        return str(int(num))
    return f"{num:.2f}".replace(".", ",")


def _join_paragraphs(parts: List[str]) -> str:
    return "\n\n".join(p.strip() for p in parts if p and str(p).strip())


def _pluralize(count: Any, singular: str, plural: str) -> Optional[str]:
    try:
        n = int(count)
    except (TypeError, ValueError):
        return None
    if n <= 0:
        return None
    return f"{n} {singular if n == 1 else plural}"


def _display_instance_name(name: Any, instance_id: Any = None) -> str:
    if instance_id and str(instance_id) in INSTANCE_DISPLAY_NAMES:
        return INSTANCE_DISPLAY_NAMES[str(instance_id)]
    raw = str(name or "").strip()
    return INSTANCE_DISPLAY_NAMES.get(raw, raw)


def _render_optionen_text(options: List[Dict[str, Any]]) -> str:
    """Fallback-Text falls die dynamische Tabelle nicht greift."""
    if not options:
        return "Keine optionalen Softwaremodule gewählt."
    blocks = []
    for idx, opt in enumerate(options, start=1):
        title = (opt.get("title") or "").strip()
        text = (opt.get("text") or "").strip()
        if title and text:
            blocks.append(f"{idx}. {title}\n{text}")
        elif title:
            blocks.append(f"{idx}. {title}")
        elif text:
            blocks.append(f"{idx}. {text}")
    return "\n\n".join(blocks)


def _set_tc_paragraphs(
    tc,
    lines: List[str],
    *,
    first_bold: bool = False,
    align_right: bool = False,
    font_size_half_points: int = 22,
) -> None:
    """Ersetzt den Zellinhalt durch Absätze (eine Zeile = ein Absatz)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for child in list(tc):
        if child.tag == qn("w:p"):
            tc.remove(child)

    if not lines:
        lines = [""]

    for i, line in enumerate(lines):
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:before"), "40")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        pPr.append(spacing)
        if align_right:
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "right")
            pPr.append(jc)
        p.append(pPr)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        if first_bold and i == 0 and str(line).strip():
            rPr.append(OxmlElement("w:b"))
            rPr.append(OxmlElement("w:bCs"))
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size_half_points)))
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(int(font_size_half_points)))
        rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = str(line)
        r.append(t)
        p.append(r)
        tc.append(p)


def _set_row_min_height(tr, twips: int = 1600) -> None:
    """Mindestzeilenhöhe für mehr Platz pro Option."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    # bestehende Höhe entfernen
    for old in trPr.findall(qn("w:trHeight")):
        trPr.remove(old)
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(twips)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def apply_options_table(
    docx_bytes: bytes,
    options: List[Dict[str, Any]],
    *,
    total_label: str = "Total Softwarelizenzen exkl. MwSt.",
    total_price: str = "",
) -> bytes:
    """
    Optionen-Tabelle: gewählte Module (Titel + Text). Preise stehen in der Zusammenfassung.
    """
    from copy import deepcopy

    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(docx_bytes))
    table = None
    for candidate in doc.tables:
        if not candidate.rows:
            continue
        header = "".join((candidate.rows[0].cells[0].text or "").split())
        if header.lower().startswith("optionen"):
            table = candidate
            break
    if table is None:
        return docx_bytes

    tbl = table._tbl
    data_rows = list(table.rows)[1:]
    if not data_rows:
        return docx_bytes

    template_tr = deepcopy(data_rows[0]._tr)
    for row in data_rows:
        tbl.remove(row._tr)

    def add_row(cell_texts: List[Any], *, min_height: int = 900, first_bold: bool = False) -> None:
        tr = deepcopy(template_tr)
        tcs = tr.findall(qn("w:tc"))
        values = list(cell_texts)
        while len(values) < len(tcs):
            values.append("")
        for i, tc in enumerate(tcs):
            raw = values[i]
            if isinstance(raw, (list, tuple)):
                lines = [str(x) for x in raw]
            else:
                lines = str(raw).split("\n") if raw is not None else [""]
            _set_tc_paragraphs(tc, lines, first_bold=(first_bold and i == 0))
        _set_row_min_height(tr, min_height)
        tbl.append(tr)

    if not options:
        add_row(
            [
                "Keine optionalen Softwaremodule gewählt. "
                "Kalkulierte Positionen und Preise stehen unter «Leistungsumfang & Preise» / Zusammenfassung.",
                "",
                "",
                "",
            ],
            min_height=900,
        )
    else:
        for idx, opt in enumerate(options, start=1):
            title = (opt.get("title") or "").strip() or f"Option {idx}"
            text = (opt.get("text") or "").strip()
            lines = [f"{idx}. {title}"]
            if text:
                lines.append(text)
            lines.append("")
            add_row([lines, "", "", ""], min_height=1400, first_bold=True)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _set_tc_text_simple(tc, text: str, *, bold: bool = True, fill: Optional[str] = None) -> None:
    """Kurzer Zelltext (z. B. gelber Tabellenkopf)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if fill:
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        for old in tcPr.findall(qn("w:shd")):
            tcPr.remove(old)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)
    _set_tc_paragraphs(tc, [text], first_bold=bold)


def apply_zusammenfassung_table(
    docx_bytes: bytes,
    scope_groups: List[Dict[str, Any]],
    *,
    subtotal_chf: Any = None,
    discount_chf: Any = None,
    grand_total_chf: Any = None,
) -> bytes:
    """
    Zwei Zusammenfassungstabellen mit gelbem Kopf:
      1) Softwarelizenzen → Total A
      2) IT-Aufwand → Total B
    Danach fett: Total netto exkl. MwSt.
    """
    from copy import deepcopy

    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(docx_bytes))
    z_table = None
    opt_table = None
    for candidate in doc.tables:
        if not candidate.rows:
            continue
        header = "".join((candidate.rows[0].cells[0].text or "").split()).lower()
        if header.startswith("zusammenfassung"):
            z_table = candidate
        elif header.startswith("optionen"):
            opt_table = candidate

    if z_table is None:
        return docx_bytes

    template_tr = None
    if opt_table is not None and len(opt_table.rows) > 1:
        template_tr = deepcopy(opt_table.rows[1]._tr)
    elif len(z_table.rows) > 1:
        best = max(z_table.rows[1:], key=lambda r: len(r._tr.findall(qn("w:tc"))))
        template_tr = deepcopy(best._tr)
    if template_tr is None:
        return docx_bytes

    header_tr_template = deepcopy(z_table.rows[0]._tr)

    def classify_groups() -> Dict[str, List[Dict[str, Any]]]:
        lic: List[Dict[str, Any]] = []
        it: List[Dict[str, Any]] = []
        for group in scope_groups or []:
            gid = str(group.get("id") or "").lower()
            title = str(group.get("title") or "").lower()
            if gid == "license" or "lizenz" in title or "software" in title:
                lic.append(group)
            elif gid == "it" or "it-" in title or "it " in title or "reise" in title:
                it.append(group)
            else:
                # Fallback: unklare Gruppen zu Lizenzen
                lic.append(group)
        return {"license": lic, "it": it}

    def _grid_col_count(table) -> int:
        grid = table._tbl.tblGrid
        if grid is not None:
            cols = grid.findall(qn("w:gridCol"))
            if cols:
                return len(cols)
        # Fallback aus Header-gridSpan / Zellen
        tcs = table.rows[0]._tr.findall(qn("w:tc")) if table.rows else []
        span = 0
        for tc in tcs:
            tcPr = tc.find(qn("w:tcPr"))
            gs = tcPr.find(qn("w:gridSpan")) if tcPr is not None else None
            span += int(gs.get(qn("w:val"))) if gs is not None else 1
        return max(span, 2)

    def _set_grid_span(tc, span: int, width_dxa: Optional[int] = None) -> None:
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        for old in tcPr.findall(qn("w:gridSpan")):
            tcPr.remove(old)
        if span > 1:
            gs = OxmlElement("w:gridSpan")
            gs.set(qn("w:val"), str(int(span)))
            tcPr.append(gs)
        if width_dxa is not None:
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(width_dxa)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)

    def _table_width_dxa(table) -> int:
        grid = table._tbl.tblGrid
        if grid is None:
            return 9000
        total = 0
        for col in grid.findall(qn("w:gridCol")):
            try:
                total += int(col.get(qn("w:w")) or 0)
            except (TypeError, ValueError):
                pass
        return total or 9000

    def clear_data_rows(table) -> None:
        tbl = table._tbl
        for row in list(table.rows)[1:]:
            tbl.remove(row._tr)

    def set_header(table, title: str) -> None:
        tcs = table.rows[0]._tr.findall(qn("w:tc"))
        if not tcs:
            return
        # Gelber Kopf als eine volle Spalte über die gesamte Tabellenbreite
        keep = tcs[0]
        for tc in tcs[1:]:
            table.rows[0]._tr.remove(tc)
        _set_grid_span(keep, _grid_col_count(table), width_dxa=_table_width_dxa(table))
        _set_tc_text_simple(keep, title, bold=True, fill="FFED00")

    def add_content_row(
        table,
        lines: List[str],
        *,
        min_height: int = 1500,
        first_bold: bool = True,
        qty: Optional[str] = None,
    ) -> None:
        """Inhaltszeile: volle Breite, oder Text | Anzahl wenn qty gesetzt."""
        tr = deepcopy(template_tr)
        tcs = tr.findall(qn("w:tc"))
        if not tcs:
            return
        cols = _grid_col_count(table)
        total_w = _table_width_dxa(table)

        if qty is None:
            keep = tcs[0]
            for tc in tcs[1:]:
                tr.remove(tc)
            _set_grid_span(keep, cols, width_dxa=total_w)
            _set_tc_paragraphs(keep, lines, first_bold=first_bold)
        else:
            # Mindestens 2 Zellen
            while len(tr.findall(qn("w:tc"))) < 2:
                tr.append(deepcopy(tcs[0]))
            tcs = tr.findall(qn("w:tc"))
            keep_label = tcs[0]
            keep_qty = tcs[-1]
            for tc in tcs[1:-1]:
                tr.remove(tc)
            tcs = tr.findall(qn("w:tc"))
            keep_label, keep_qty = tcs[0], tcs[-1]
            qty_w = max(900, total_w // 8)
            label_w = max(2000, total_w - qty_w)
            label_span = max(1, cols - 1)
            _set_grid_span(keep_label, label_span, width_dxa=label_w)
            _set_grid_span(keep_qty, 1, width_dxa=qty_w)
            _set_tc_paragraphs(keep_label, lines, first_bold=first_bold)
            _set_tc_paragraphs(keep_qty, [str(qty)], first_bold=True, align_right=True)

        _set_row_min_height(tr, min_height)
        table._tbl.append(tr)

    def add_total_row(
        table,
        label: str,
        price: str,
        *,
        min_height: int = 750,
    ) -> None:
        """Totalzeile: 2 Spalten — Text | Preis."""
        tr = deepcopy(template_tr)
        tcs = tr.findall(qn("w:tc"))
        if len(tcs) < 2:
            # Mindestens 2 Zellen sicherstellen
            while len(tr.findall(qn("w:tc"))) < 2:
                tr.append(deepcopy(tcs[0] if tcs else template_tr.findall(qn("w:tc"))[0]))
            tcs = tr.findall(qn("w:tc"))
        # Nur erste und letzte Zelle behalten
        keep_label = tcs[0]
        keep_price = tcs[-1]
        for tc in tcs[1:-1]:
            tr.remove(tc)
        # Nach Entfernen der Mittelzellen Reihenfolge prüfen
        tcs = tr.findall(qn("w:tc"))
        if len(tcs) == 1:
            # template hatte 1 Zelle — Preis-Zelle klonen
            keep_price = deepcopy(keep_label)
            tr.append(keep_price)
            tcs = tr.findall(qn("w:tc"))
        keep_label, keep_price = tcs[0], tcs[-1]

        cols = _grid_col_count(table)
        total_w = _table_width_dxa(table)
        price_w = max(1800, total_w // 4)
        label_w = max(2000, total_w - price_w)
        label_span = max(1, cols - 1)
        _set_grid_span(keep_label, label_span, width_dxa=label_w)
        _set_grid_span(keep_price, 1, width_dxa=price_w)
        _set_tc_paragraphs(keep_label, [label], first_bold=True)
        _set_tc_paragraphs(keep_price, [price], first_bold=True, align_right=True)
        _set_row_min_height(tr, min_height)
        table._tbl.append(tr)

    def fill_group_table(
        table,
        groups: List[Dict[str, Any]],
        *,
        total_prefix: str,
        header_title: str,
        show_qty: bool = False,
    ) -> None:
        set_header(table, header_title)
        clear_data_rows(table)
        pos = 0
        if not groups:
            add_content_row(table, ["Keine Positionen in diesem Bereich."], min_height=800, first_bold=False)
            return
        if show_qty:
            # Spaltenköpfe: Position | Anzahl
            add_content_row(table, ["Position"], min_height=420, first_bold=True, qty="Anzahl")
        for group in groups:
            for item in group.get("items") or []:
                pos += 1
                name = (item.get("name") or "").strip() or f"Position {pos}"
                desc = (item.get("description") or "").strip()
                lines = [f"{pos}. {name}"]
                if desc:
                    lines.append(desc)
                lines.append("")
                qty_val = item.get("qty")
                qty_text = None
                if show_qty:
                    if qty_val is None or qty_val == "":
                        qty_text = "—"
                    else:
                        try:
                            qn_val = float(qty_val)
                            qty_text = str(int(qn_val)) if qn_val == int(qn_val) else str(qn_val)
                        except (TypeError, ValueError):
                            qty_text = str(qty_val)
                add_content_row(
                    table,
                    lines,
                    min_height=1500,
                    first_bold=True,
                    qty=qty_text,
                )
            total = group.get("total")
            if total_prefix == "A":
                label = "Total A · Softwarelizenzen"
            elif total_prefix == "B":
                label = "Total B · IT-Aufwand"
            else:
                label = f"Total {(group.get('title') or '').strip()}".strip() or "Total"
            if total is not None:
                add_total_row(
                    table,
                    label,
                    _money(total, group.get("currency") or "CHF"),
                    min_height=750,
                )

    buckets = classify_groups()
    # Tabelle 1: bestehende Zusammenfassung → Softwarelizenzen
    fill_group_table(
        z_table,
        buckets["license"],
        total_prefix="A",
        header_title="Zusammenfassung – Softwarelizenzen",
        show_qty=True,
    )

    # Tabelle 2: Kopie mit gleichem gelben Kopf → IT-Aufwand
    z_tbl = z_table._tbl
    it_tbl = deepcopy(z_tbl)
    # Nach der ersten Zusammenfassung einfügen (+ Leerabsatz dazwischen)
    parent = z_tbl.getparent()
    insert_at = list(parent).index(z_tbl) + 1
    spacer = OxmlElement("w:p")
    parent.insert(insert_at, spacer)
    parent.insert(insert_at + 1, it_tbl)

    # python-docx Table-Wrapper für die Kopie
    from docx.table import Table

    it_table = Table(it_tbl, z_table._parent)
    # Headerzeile der Kopie ggf. ersetzen falls beim Clear beschädigt
    if not it_table.rows:
        it_tbl.append(deepcopy(header_tr_template))
        it_table = Table(it_tbl, z_table._parent)
    fill_group_table(
        it_table,
        buckets["it"],
        total_prefix="B",
        header_title="Zusammenfassung – IT-Aufwand",
    )

    # Gesamttotal als eigene Mini-Tabelle (gelb/fett hervorgehoben)
    grand_tbl = deepcopy(z_tbl)
    # Nur Kopf + eine Datenzeile behalten
    # Zuerst leeren und neu füllen über Table-API
    parent.insert(insert_at + 2, OxmlElement("w:p"))
    parent.insert(insert_at + 3, grand_tbl)
    grand_table = Table(grand_tbl, z_table._parent)
    set_header(grand_table, "Gesamttotal")
    clear_data_rows(grand_table)

    if subtotal_chf is not None and discount_chf:
        add_total_row(grand_table, "Zwischentotal", _money(subtotal_chf, "CHF"), min_height=600)
        add_total_row(grand_table, "Projektrabatt", f"− {_money(discount_chf, 'CHF')}", min_height=600)

    if grand_total_chf is not None:
        add_total_row(
            grand_table,
            "Total netto exkl. MwSt.",
            _money(grand_total_chf, "CHF"),
            min_height=900,
        )
    else:
        add_total_row(grand_table, "Total netto exkl. MwSt.", "—", min_height=900)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _render_section_body(sec: Dict[str, Any], blocks: List[str], heading_prefix: str = "") -> None:
    sid = (sec.get("id") or "").strip()
    title = (sec.get("title") or "").strip()
    head = f"{heading_prefix}{sid} {title}".strip()
    if head:
        blocks.append(head)
    for p in sec.get("paragraphs") or []:
        if p:
            blocks.append(str(p).strip())
    for b in sec.get("bullets") or []:
        if b:
            blocks.append(f"• {str(b).strip()}")
    table = sec.get("table")
    if isinstance(table, dict):
        headers = [str(h).strip() for h in (table.get("headers") or []) if str(h).strip()]
        if headers:
            blocks.append(" | ".join(headers))
        for row in table.get("rows") or []:
            cells = [str(c).strip() for c in row]
            if any(cells):
                blocks.append(" | ".join(cells))
    for p in sec.get("paragraphsAfter") or []:
        if p:
            blocks.append(str(p).strip())
    for sub in sec.get("subsections") or []:
        if isinstance(sub, dict):
            _render_section_body(sub, blocks)


def _render_bedingungen_text(terms: Dict[str, Any]) -> str:
    blocks: List[str] = []
    intro = (terms.get("intro") or "").strip()
    if intro:
        blocks.append(intro)
    for sec in terms.get("sections") or []:
        if isinstance(sec, dict):
            _render_section_body(sec, blocks)
    return _join_paragraphs(blocks) if blocks else "—"


def _greeting_line(contact: str) -> str:
    name = (contact or "").strip()
    if not name:
        return "Sehr geehrte Damen und Herren,"
    tokens = name.split()
    first = tokens[0].lower().rstrip(".")
    if first in {"herr", "hr"}:
        return f"Sehr geehrter {' '.join(tokens)},"
    if first in {"frau", "fr"}:
        return f"Sehr geehrte {' '.join(tokens)},"
    return f"Sehr geehrte/r {name},"


def _set_paragraph_text(paragraph, text: str) -> None:
    """Ersetzt den gesamten Absatztext (behält grob die Formatierung des ersten Runs)."""
    value = text or ""
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def _insert_page_break_before(element) -> None:
    """Fügt vor einem Body-Element (Absatz/Tabelle) einen Seitenumbruch ein."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    element.addprevious(p)


def _set_paragraph_lines(paragraph, lines: List[str]) -> None:
    """Setzt Absatz mit weichen Zeilenumbrüchen (Titel / Projektnr. / Datum)."""
    from docx.oxml import OxmlElement

    clean = [str(x).strip() for x in lines if str(x or "").strip()]
    if not clean:
        _set_paragraph_text(paragraph, "")
        return
    # Clear existing runs
    if paragraph.runs:
        paragraph.runs[0].text = clean[0]
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(clean[0])
    for line in clean[1:]:
        run = paragraph.add_run()
        br = OxmlElement("w:br")
        run._r.append(br)
        paragraph.add_run(line)


def polish_offer_docx(docx_bytes: bytes, context: Dict[str, Any]) -> bytes:
    """Nachbearbeitung: Cover/Index, Anrede, Versionstabelle, Seitenumbrüche."""
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    angebotsnummer = str(context.get("angebotsnummer") or "").strip()
    revision_code = str(context.get("revision_code") or "").strip()
    datum = str(context.get("datum") or "").strip()
    titel = str(context.get("titel") or "").strip()
    anrede = str(context.get("anrede") or "").strip()
    einleitung = str(context.get("einleitung") or "").strip()
    erstellt_von = str(context.get("erstellt_von") or "").strip()
    project_no = angebotsnummer or revision_code

    # Cover: Titel · Projektnr. · Datum als getrennte Zeilen (nicht in einer Headline)
    for p in doc.paragraphs[:8]:
        t = (p.text or "").strip()
        if "WAMAS" in t and ("Angebot" in t or "Lift" in t):
            lines = [titel or "Angebot WAMAS® Lift & Store"]
            if project_no:
                lines.append(project_no)
            if datum:
                lines.append(datum)
            _set_paragraph_lines(p, lines)
            break

    # Anrede dynamisch
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("Sehr geehrter Herr") or t.startswith("Sehr geehrte"):
            _set_paragraph_text(p, anrede or "Sehr geehrte Damen und Herren,")
            break

    # Veralteten Hardcode-Dankestext durch Einleitung ersetzen
    for p in doc.paragraphs:
        t = p.text or ""
        if "14.04.2008" in t or "Wir danken für Ihre Anfrage vom" in t:
            _set_paragraph_text(
                p,
                einleitung
                or "Wir danken für Ihre Anfrage und unterbreiten Ihnen nachstehend unser Angebot.",
            )
            break

    # Bedingungen-Versionszeile (falls vorhanden)
    terms_version = str(context.get("terms_version") or "").strip()
    if terms_version:
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t.startswith("Version ") and len(t) < 40 and any(ch.isdigit() for ch in t):
                _set_paragraph_text(p, f"Version {terms_version}")
                break

    # Versionstabelle: Index | Datum | Erstellt von | Kommentar/Angebotsnr.
    for table in doc.tables:
        if not table.rows:
            continue
        header = (table.rows[0].cells[0].text or "").strip().lower()
        if not header.startswith("version"):
            continue
        if len(table.rows) < 2 or len(table.rows[1].cells) < 3:
            continue
        row = table.rows[1]
        row.cells[0].text = revision_code or angebotsnummer or ""
        if len(row.cells) > 1:
            row.cells[1].text = datum
        if len(row.cells) > 2 and not (row.cells[2].text or "").strip():
            row.cells[2].text = erstellt_von
        if len(row.cells) > 3:
            comment = angebotsnummer if angebotsnummer != revision_code else ""
            if revision_code and angebotsnummer and revision_code not in angebotsnummer:
                comment = angebotsnummer
            row.cells[3].text = comment
        break

    # Seitenumbrüche vor zentralen Blöcken (analog HTML-Druck)
    break_headers = (
        "einleitung",
        "zusammenfassung",
        "kaufmännische bedingungen",
        "kaufmaennische bedingungen",
        "vertragsbedingungen",
        "zuständigkeiten",
        "zustaendigkeiten",
        "a4.",
        "a4 ",
    )
    seen = set()
    for table in doc.tables:
        if not table.rows:
            continue
        header = (table.rows[0].cells[0].text or "").strip().lower()
        key = next((h for h in break_headers if header.startswith(h) or h in header), None)
        if not key or key in seen:
            continue
        seen.add(key)
        _insert_page_break_before(table._tbl)

    for p in doc.paragraphs:
        t = (p.text or "").strip().lower()
        if not t:
            continue
        key = None
        if t.startswith("a4") or "zuständigkeit" in t or "zustaendigkeit" in t:
            key = "zuständigkeiten"
        elif t.startswith("vertragsbedingungen") or t.startswith("kaufmännische bedingungen"):
            key = "vertragsbedingungen"
        if not key or key in seen:
            continue
        seen.add(key)
        _insert_page_break_before(p._p)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _split_swiss_address(address: str) -> Dict[str, str]:
    """Split 'Strasse Nr, PLZ Ort' into cover fields."""
    raw = re.sub(r"\s+", " ", (address or "").strip())
    if not raw:
        return {"strasse": "", "plz_ort": "", "adresse": ""}
    match = re.match(r"^(.*?),\s*(\d{4}\s+.+)$", raw)
    if match:
        return {
            "strasse": match.group(1).strip(),
            "plz_ort": match.group(2).strip(),
            "adresse": raw,
        }
    # Fallback: last token group with PLZ
    match = re.match(r"^(.*?)(\d{4}\s+.+)$", raw)
    if match and match.group(1).strip():
        return {
            "strasse": match.group(1).strip(" ,"),
            "plz_ort": match.group(2).strip(),
            "adresse": raw,
        }
    return {"strasse": raw, "plz_ort": "", "adresse": raw}


def build_template_context(offer: Dict[str, Any]) -> Dict[str, Any]:
    """Mappt Angebots-JSON auf die Platzhalter-Felder der Word-Vorlage."""
    meta = offer.get("meta") or {}
    customer = offer.get("customer") or {}
    content = offer.get("content") or {}
    cfg = content.get("configurationSummary") or {}
    summary = offer.get("priceSummary") or {}
    lic = summary.get("license") or {}
    it = summary.get("it") or {}
    terms = content.get("commercialTerms") or {}
    closing = terms.get("closing") or {}
    signatories = closing.get("signatories") or [{}, {}]
    sig1 = signatories[0] if len(signatories) > 0 else {}
    sig2 = signatories[1] if len(signatories) > 1 else {}

    instance_name = _display_instance_name(cfg.get("instanceName"), cfg.get("instanceId"))
    cfg_bits = " · ".join(
        x
        for x in [
            instance_name or None,
            f"{cfg['instanceCount']}×" if cfg.get("instanceCount") else None,
            _pluralize(cfg.get("deviceCount"), "Gerät", "Geräte"),
            _pluralize(cfg.get("zoneCount"), "Zone", "Zonen"),
            _pluralize(cfg.get("openingCount"), "Öffnung", "Öffnungen"),
        ]
        if x
    )

    # Kunden-Word: Leistungsumfang ohne Positionspreise/Stunden
    lines: List[Dict[str, Any]] = []
    scope_groups = offer.get("scopeGroups") or content.get("scopeGroups") or []
    if scope_groups:
        pos = 0
        for group in scope_groups:
            for item in group.get("items") or []:
                pos += 1
                lines.append(
                    {
                        "pos": pos,
                        "section": group.get("title") or "",
                        "sku": "",
                        "name": item.get("name") or "",
                        "description": item.get("description") or "",
                        "qty": item.get("qty") if item.get("qty") is not None else "",
                        "unit_price": "",
                        "hours": "",
                        "amount": "",
                        "currency": "",
                    }
                )
    else:
        for line in offer.get("commercialLines") or []:
            desc = str(line.get("description") or "")
            desc = re.sub(r"\s*·\s*IC\s+[−\-]?\s*[\d.'\s,]+\s*EUR", "", desc, flags=re.IGNORECASE)
            desc = re.sub(r"\s*IC\s+[−\-]?\s*[\d.'\s,]+\s*EUR", "", desc, flags=re.IGNORECASE).strip()
            lines.append(
                {
                    "pos": line.get("pos", ""),
                    "section": line.get("section", ""),
                    "sku": line.get("sku", ""),
                    "name": line.get("name", ""),
                    "description": desc,
                    "qty": "",
                    "unit_price": "",
                    "hours": "",
                    "amount": "",
                    "currency": "",
                }
            )

    # Keine Bereichs-Totals hier: Projektrabatt gilt nur für Gesamttotal.
    scope_text_blocks = []
    for group in scope_groups:
        scope_text_blocks.append(group.get("title") or "")
        for item in group.get("items") or []:
            name = (item.get("name") or "").strip()
            desc = (item.get("description") or "").strip()
            if name and desc:
                scope_text_blocks.append(f"• {name}\n{desc}")
            elif name:
                scope_text_blocks.append(f"• {name}")
    scope_text = _join_paragraphs(scope_text_blocks)

    einleitung = _join_paragraphs(
        [
            content.get("intro") or "",
            content.get("introVariant") or "",
            content.get("recommendation") or "",
        ]
    )

    angebotsnummer = meta.get("archiveTitle") or meta.get("offerNumber", "")
    datum = meta.get("documentDate", "")
    gueltig_bis = meta.get("validUntil", "")
    revision_von = (
        meta.get("revisionOf")
        or (offer.get("sources") or {}).get("basedOnOfferNumber")
        or ""
    )
    meta_parts = [
        f"Angebot {angebotsnummer}" if angebotsnummer else "Angebot",
        datum,
        f"Gültig bis {gueltig_bis}" if gueltig_bis else "",
        f"Rev. von {revision_von}" if revision_von else "",
    ]
    meta_zeile = "  ·  ".join(p for p in meta_parts if p)

    note = summary.get("note")
    if note and "IC-Preise" in str(note):
        note = "Alle Preise in CHF, exkl. MwSt."
    if not note:
        note = "Alle Preise in CHF, exkl. MwSt."

    addr = _split_swiss_address(customer.get("address", ""))
    prepared_by = (meta.get("preparedBy") or "").strip()
    # Cover: SSI-Ansprechpartner | Letzte Seite: Unterschriften (getrennt)
    ssi_contacts = offer.get("ssiContacts") or meta.get("ssiContacts") or []
    ssi1 = dict(ssi_contacts[0] if len(ssi_contacts) > 0 else {})
    ssi2 = dict(ssi_contacts[1] if len(ssi_contacts) > 1 else {})
    # Signaturen aus Bedingungen / Angebot — nicht dieselben wie Cover
    offer_sigs = offer.get("signatories") or []
    raw_sig1 = dict(offer_sigs[0] if len(offer_sigs) > 0 else sig1 or {})
    raw_sig2 = dict(offer_sigs[1] if len(offer_sigs) > 1 else sig2 or {})

    def _enrich(contact: Dict[str, Any], *, name_hint: str = "") -> Dict[str, Any]:
        if contact.get("email") and contact.get("name"):
            return contact
        found = _lookup_ssi_contact(
            contact_id=str(contact.get("id") or ""),
            name=str(contact.get("name") or name_hint or ""),
        )
        if found:
            merged = dict(found)
            merged.update({k: v for k, v in contact.items() if v})
            return merged
        return contact

    ssi1 = _enrich(ssi1, name_hint=prepared_by)
    ssi2 = _enrich(ssi2)
    raw_sig1 = _enrich(raw_sig1)
    raw_sig2 = _enrich(raw_sig2)

    def _fields(primary: Dict[str, Any], *, name_fallback: str = "") -> Dict[str, str]:
        """Nur Felder der gewählten Person — kein Fallback auf andere Signatur."""
        name = (primary.get("name") or name_fallback or "").strip()
        if not name and not primary.get("id"):
            return {"name": "", "title": "", "role": "", "email": "", "phone": ""}
        return {
            "name": name,
            "title": (primary.get("title") or "").strip(),
            "role": (primary.get("role") or "").strip(),
            "email": (primary.get("email") or "").strip(),
            "phone": (primary.get("phone") or "").strip(),
        }

    f1 = _fields(ssi1, name_fallback=prepared_by)
    f2 = _fields(ssi2)
    sf1 = _fields(raw_sig1)
    sf2 = _fields(raw_sig2)
    # Cover-Fallback nur wenn leer
    if not f1["name"] and not f2["name"]:
        f1 = _fields(sig1)
        f2 = _fields(sig2)
    # Signatur-Fallback auf Cover nur wenn Signaturen komplett leer
    if not sf1["name"] and not sf2["name"]:
        sf1, sf2 = f1, f2

    def _position(title: str, role: str = "") -> str:
        title = (title or "").strip()
        role = (role or "").strip()
        if title and role:
            return f"{title} / {role}"
        return title or role

    ssi1_name = f1["name"]
    ssi1_pos = _position(f1["title"], f1["role"])
    ssi1_mail, ssi1_tel = f1["email"], f1["phone"]
    ssi2_name = f2["name"]
    ssi2_pos = _position(f2["title"], f2["role"])
    ssi2_mail, ssi2_tel = f2["email"], f2["phone"]
    sig1_name, sig1_title, sig1_role = sf1["name"], sf1["title"], sf1["role"]
    sig2_name, sig2_title, sig2_role = sf2["name"], sf2["title"], sf2["role"]
    sig1_mail, sig2_mail = sf1["email"], sf2["email"]
    if prepared_by == "" and ssi1_name:
        prepared_by = ssi1_name

    def _details(pos: str, email: str, phone: str) -> str:
        parts = [p for p in [pos, email, phone] if p]
        return "\n".join(parts)

    def _sig_details(title: str, role: str, email: str = "") -> str:
        parts = [p for p in [title, role, email] if p]
        return "\n".join(parts)

    return {
        "dokument_label": content.get("documentLabel") or "Angebot / Preisliste",
        "titel": content.get("title") or "Angebot WAMAS® Lift & Store",
        "untertitel": content.get("subtitle")
        or "SSI SCHÄFER · Softwarelösung für Vertical Lift Modules (SSI LOGIMAT®)",
        "angebotsnummer": angebotsnummer,
        "datum": datum,
        "ort_datum": (
            f"Schaffhauserstrasse 10, 8213 Neunkirch · {datum}"
            if datum
            else "Schaffhauserstrasse 10, 8213 Neunkirch"
        ),
        "gueltig_bis": gueltig_bis,
        "meta_zeile": meta_zeile,
        "revision_code": meta.get("revisionCode") or "",
        "version_software": (offer.get("branding") or {}).get("version", "2.8"),
        "erstellt_von": prepared_by,
        "revision_von": revision_von,
        "kunde": customer.get("company", ""),
        "projekt": customer.get("projectName", ""),
        "ansprechpartner": customer.get("contact", ""),
        "anrede": _greeting_line(customer.get("contact", "")),
        "email": customer.get("email", ""),
        "telefon": customer.get("phone", ""),
        "fax": customer.get("fax", "") or "",
        "adresse": addr["adresse"],
        "strasse": addr["strasse"],
        "plz_ort": addr["plz_ort"],
        "ssi_kontakt_1_name": ssi1_name,
        "ssi_kontakt_1_position": ssi1_pos,
        "ssi_kontakt_1_email": ssi1_mail,
        "ssi_kontakt_1_telefon": ssi1_tel,
        "ssi_kontakt_1_details": _details(ssi1_pos, ssi1_mail, ssi1_tel),
        "ssi_kontakt_2_name": ssi2_name,
        "ssi_kontakt_2_position": ssi2_pos,
        "ssi_kontakt_2_email": ssi2_mail,
        "ssi_kontakt_2_telefon": ssi2_tel,
        "ssi_kontakt_2_details": _details(ssi2_pos, ssi2_mail, ssi2_tel),
        "konfiguration": cfg_bits,
        "instance_name": instance_name,
        "instance_count": cfg.get("instanceCount") or "",
        "device_count": cfg.get("deviceCount") or "",
        "zone_count": cfg.get("zoneCount") or "",
        "opening_count": cfg.get("openingCount") or "",
        "order_handling": "Ja" if cfg.get("hasOrderHandling") else "Nein",
        "einleitung": einleitung,
        "optionen_text": _render_optionen_text(content.get("selectedOptions") or []),
        "bedingungen_text": _render_bedingungen_text(terms),
        "leistungsumfang_text": scope_text,
        "preis_hinweis": note,
        "lizenz_total_chf": _money(lic.get("total"), lic.get("currency") or "CHF") if lic else "—",
        "lizenz_ic_eur": "",
        "marge_percent": "",
        "kurs_eur_chf": "",
        "it_total_chf": _money(it.get("total"), it.get("currency") or "CHF") if it else "—",
        "it_work_chf": "",
        "it_travel_chf": "",
        "gesamt_chf": _money(summary.get("grandTotalChf"), "CHF")
        if summary.get("grandTotalChf") is not None
        else "—",
        "lines": lines,
        "terms_version": terms.get("version", "09.2022"),
        "terms_validity_days": terms.get("validityDays", 14),
        "schluss_text": closing.get("text")
        or "Ihrem Auftrag, dem wir unsere volle Aufmerksamkeit schenken, sehen wir gerne entgegen.",
        "gruss": closing.get("greeting") or "Freundliche Grüsse",
        "firma_ssi": closing.get("company") or "SSI SCHÄFER AG",
        "signatur_1_name": sig1_name,
        "signatur_1_titel": sig1_title,
        "signatur_1_rolle": sig1_role,
        "signatur_1_details": _sig_details(sig1_title, sig1_role, sig1_mail),
        "signatur_2_name": sig2_name,
        "signatur_2_titel": sig2_title,
        "signatur_2_rolle": sig2_role,
        "signatur_2_details": _sig_details(sig2_title, sig2_role, sig2_mail),
    }


def _set_sdt_plain_text(sdt_xml: str, text: str) -> str:
    """Ersetzt den sichtbaren Inhalt eines w:sdt und entfernt Platzhalter-Flag."""
    safe = xml_escape(text or "")
    lines = safe.split("\n")
    if len(lines) <= 1:
        run_inner = f"<w:t>{safe}</w:t>"
    else:
        parts = [f'<w:t xml:space="preserve">{lines[0]}</w:t>']
        for line in lines[1:]:
            parts.append("<w:br/>")
            parts.append(f'<w:t xml:space="preserve">{line}</w:t>')
        run_inner = "".join(parts)

    def repl_content(match: re.Match) -> str:
        block = match.group(0)
        replaced = False

        def repl_run(rm: re.Match) -> str:
            nonlocal replaced
            if replaced:
                return ""
            replaced = True
            # Keep run props if present
            open_tag = rm.group(1)
            return f"{open_tag}{run_inner}</w:r>"

        block = re.sub(r"(<w:r\b[^>]*>).*?</w:r>", repl_run, block, flags=re.DOTALL)
        if not replaced:
            block = re.sub(
                r"(<w:sdtContent\b[^>]*>)(.*?)(</w:sdtContent>)",
                rf"\1<w:p><w:r>{run_inner}</w:r></w:p>\3",
                block,
                count=1,
                flags=re.DOTALL,
            )
        return block

    out = re.sub(
        r"<w:sdtContent\b[^>]*>.*?</w:sdtContent>",
        repl_content,
        sdt_xml,
        count=1,
        flags=re.DOTALL,
    )
    out = re.sub(r"<w:showingPlcHdr\s*/>", "", out)
    out = re.sub(r'<w:rStyle\s+w:val="Platzhaltertext"\s*/>', "", out)
    return out


def _iter_top_level_sdts(xml: str):
    """Yield (start, end, chunk) for each top-level w:sdt element."""
    i = 0
    while True:
        start = xml.find("<w:sdt>", i)
        if start < 0:
            return
        depth = 0
        j = start
        while j < len(xml):
            if xml.startswith("<w:sdt>", j):
                depth += 1
                j += 7
                continue
            if xml.startswith("</w:sdt>", j):
                depth -= 1
                j += 8
                if depth == 0:
                    yield start, j, xml[start:j]
                    i = j
                    break
                continue
            j += 1
        else:
            return


def apply_sdt_contact_fields(docx_bytes: bytes, context: Dict[str, Any]) -> bytes:
    """Befüllt Word-Dropdowns/Textfelder (Cover + Unterschriften) nach docxtpl-Render."""
    zin = zipfile.ZipFile(io.BytesIO(docx_bytes), "r")
    try:
        xml = zin.read("word/document.xml").decode("utf-8")
    except KeyError:
        zin.close()
        return docx_bytes

    replacements = 0
    pieces: List[str] = []
    cursor = 0
    for start, end, chunk in _iter_top_level_sdts(xml):
        pieces.append(xml[cursor:start])
        alias_m = re.search(r'<w:alias\b[^>]*w:val="([^"]+)"', chunk)
        tag_m = re.search(r'<w:tag\b[^>]*w:val="([^"]+)"', chunk)
        key = (alias_m.group(1) if alias_m else "") or (tag_m.group(1) if tag_m else "")
        ctx_key = _SDT_TEXT_MAP.get(key)
        if ctx_key and context.get(ctx_key) is not None:
            pieces.append(_set_sdt_plain_text(chunk, str(context.get(ctx_key) or "")))
            replacements += 1
        else:
            pieces.append(chunk)
        cursor = end
    pieces.append(xml[cursor:])
    new_xml = "".join(pieces)

    if replacements == 0:
        zin.close()
        return docx_bytes

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                data = new_xml.encode("utf-8")
            zout.writestr(item, data)
    zin.close()
    return out.getvalue()


def build_offer_docx(offer: Dict[str, Any]) -> bytes:
    """Befüllt die Anhang-Vorlage (ein Dokument, SSI-Anhang inkl. Platzhalter)."""
    if offer.get("kind") != "offer_document":
        raise ValueError("Word-Export ist nur für Gesamtangebote (offer_document) verfügbar.")

    template_path = ensure_annex_template()
    context = build_template_context(offer)
    selected_options = list((offer.get("content") or {}).get("selectedOptions") or [])
    # optionen_text bleibt die aus dem Context gerenderte Optionsbeschreibung

    tpl = DocxTemplate(str(template_path))
    tpl.render(context, autoescape=True)

    out = io.BytesIO()
    tpl.save(out)
    docx_bytes = apply_sdt_contact_fields(out.getvalue(), context)
    docx_bytes = apply_options_table(docx_bytes, selected_options)
    summary = offer.get("priceSummary") or {}
    scope_groups = offer.get("scopeGroups") or (offer.get("content") or {}).get("scopeGroups") or []
    docx_bytes = apply_zusammenfassung_table(
        docx_bytes,
        list(scope_groups),
        subtotal_chf=summary.get("subtotalChf"),
        discount_chf=summary.get("commercialDiscountChf"),
        grand_total_chf=summary.get("grandTotalChf"),
    )
    return polish_offer_docx(docx_bytes, context)

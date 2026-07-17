#!/usr/bin/env python3
"""Erzeugt die Platzhalter-Vorlage angebot_platzhalter.docx (docxtpl).

Komplexe Abschnitte (Optionen, Bedingungen) werden als vorbereitete
Textfelder {{optionen_text}} / {{bedingungen_text}} gefüllt — so bleiben
Jinja-Tags im Word-XML stabil. Nur die Preistabelle nutzt {%tr for%}.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "templates" / "angebot_platzhalter.docx"

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
TEAL = RGBColor(0x00, 0x7A, 0x8A)
GRAY = RGBColor(0x4A, 0x55, 0x68)


def set_run(run, *, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_p(doc, text, *, size=11, bold=False, color=None, space_after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return p


def add_h(doc, text, size=14):
    return add_p(doc, text, size=size, bold=True, color=NAVY, space_after=8)


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

    add_p(doc, "{{dokument_label}}", size=10, bold=True, color=TEAL, space_after=2)
    add_p(doc, "{{titel}}", size=18, bold=True, color=NAVY, space_after=4)
    add_p(doc, "{{untertitel}}", size=11, color=TEAL, space_after=10)
    add_p(doc, "{{meta_zeile}}", size=10, color=GRAY, space_after=14)

    add_h(doc, "1. Projekt & Kunde")
    meta = [
        ("Kunde", "{{kunde}}"),
        ("Projekt", "{{projekt}}"),
        ("Adresse", "{{adresse}}"),
        ("Ansprechpartner", "{{ansprechpartner}}"),
        ("E-Mail", "{{email}}"),
        ("Erstellt von", "{{erstellt_von}}"),
        ("Software-Version", "{{version_software}}"),
        ("Konfiguration", "{{konfiguration}}"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    for i, (label, ph) in enumerate(meta):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = ph
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run(
                        r,
                        size=10,
                        bold=(cell is table.rows[i].cells[0]),
                        color=NAVY if cell is table.rows[i].cells[0] else None,
                    )

    add_p(doc, "", space_after=8)
    add_h(doc, "2. Einleitung")
    add_p(doc, "{{einleitung}}", size=10, space_after=10)

    add_h(doc, "3. Optionale Leistungen")
    add_p(doc, "{{optionen_text}}", size=10, space_after=10)

    add_h(doc, "4. Preiszusammenstellung")
    add_p(doc, "{{preis_hinweis}}", size=9, color=GRAY, space_after=6)

    # docxtpl: {%tr for%} und {%tr endfor%} je in eigener Zeile;
    # die mittlere Zeile ist die wiederholte Datenzeile.
    pt = doc.add_table(rows=4, cols=6)
    pt.style = "Table Grid"
    headers = ["Pos", "Bezeichnung", "Menge", "EP", "Std.", "Betrag"]
    for i, h in enumerate(headers):
        pt.rows[0].cells[i].text = h
        for p in pt.rows[0].cells[i].paragraphs:
            for r in p.runs:
                set_run(r, size=9, bold=True, color=NAVY)
    pt.rows[1].cells[0].text = "{%tr for line in lines %}"
    pt.rows[2].cells[0].text = "{{ line.pos }}"
    pt.rows[2].cells[1].text = "{{ line.name }}"
    pt.rows[2].cells[2].text = "{{ line.qty }}"
    pt.rows[2].cells[3].text = "{{ line.unit_price }}"
    pt.rows[2].cells[4].text = "{{ line.hours }}"
    pt.rows[2].cells[5].text = "{{ line.amount }} {{ line.currency }}"
    pt.rows[3].cells[0].text = "{%tr endfor %}"
    for row_idx in (1, 2, 3):
        for cell in pt.rows[row_idx].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run(r, size=9)

    add_p(doc, "", space_after=6)
    totals = [
        ("Softwarelizenzen Verkauf", "{{lizenz_total_chf}}"),
        ("IT-Aufwand / Installation", "{{it_total_chf}}"),
        ("Gesamttotal (exkl. MwSt.)", "{{gesamt_chf}}"),
    ]
    tt = doc.add_table(rows=len(totals), cols=2)
    tt.style = "Table Grid"
    for i, (label, ph) in enumerate(totals):
        tt.rows[i].cells[0].text = label
        tt.rows[i].cells[1].text = ph
        for cell in tt.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run(
                        r,
                        size=10,
                        bold=True if i == len(totals) - 1 else False,
                        color=NAVY,
                    )

    add_p(doc, "", space_after=8)
    add_h(doc, "5. Kaufmännische Bedingungen (CH {{terms_version}})")
    add_p(doc, "{{bedingungen_text}}", size=9, space_after=10)

    add_p(doc, "{{schluss_text}}", size=10, space_after=6)
    add_p(doc, "{{gruss}}", size=10, space_after=4)
    add_p(doc, "{{firma_ssi}}", size=10, bold=True, color=NAVY, space_after=10)

    add_h(doc, "Unterschriften")
    sig = doc.add_table(rows=5, cols=2)
    sig.style = "Table Grid"
    sig.rows[0].cells[0].text = "{{signatur_1_name}}"
    sig.rows[0].cells[1].text = "{{signatur_2_name}}"
    sig.rows[1].cells[0].text = "{{signatur_1_titel}}"
    sig.rows[1].cells[1].text = "{{signatur_2_titel}}"
    sig.rows[2].cells[0].text = "{{signatur_1_rolle}}"
    sig.rows[2].cells[1].text = ""
    sig.rows[3].cells[0].text = "______________________________"
    sig.rows[3].cells[1].text = "______________________________"
    sig.rows[4].cells[0].text = "SSI Schäfer AG"
    sig.rows[4].cells[1].text = "Kunde / Auftraggeber"
    for row in sig.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run(r, size=9)

    add_p(doc, "", space_after=10)
    add_p(
        doc,
        "Ab hier folgt der Software-Anhang (Funktionsbeschreibung) als Originaldokument.",
        size=9,
        color=GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()

"""Word-Export im Layout der HTML-Angebotsvorschau (programmatisch).

Die klassische Logimat-Vorlage bleibt über layout=\"template\" verfügbar.
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from docx_export import _money, _money_compact

BASE_DIR = Path(__file__).resolve().parent
STATIC_DIR = BASE_DIR / "static"
ASSETS_DIR = STATIC_DIR / "assets"

SSI_YELLOW = RGBColor(0xFF, 0xED, 0x00)
SSI_NAVY = RGBColor(0x00, 0x33, 0x66)
SSI_INK = RGBColor(0x1A, 0x1A, 0x1A)
SSI_MUTED = RGBColor(0x5A, 0x5A, 0x5A)


def _set_run_font(run, *, size_pt: float = 10, bold: bool = False, color: RGBColor = SSI_INK) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")


def _add_para(
    container,
    text: str = "",
    *,
    size_pt: float = 10,
    bold: bool = False,
    color: RGBColor = SSI_INK,
    space_after: float = 6,
    space_before: float = 0,
    align: Optional[Any] = None,
):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if text:
        run = p.add_run(text)
        _set_run_font(run, size_pt=size_pt, bold=bold, color=color)
    return p


def _shade_cell(cell, fill: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size_pt: float = 9,
    align_right: bool = False,
    color: RGBColor = SSI_INK,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align_right:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    _set_run_font(run, size_pt=size_pt, bold=bold, color=color)


def _set_cell_paragraphs(
    cell,
    lines: List[str],
    *,
    first_bold: bool = True,
    align_right: bool = False,
    size_pt: float = 9,
) -> None:
    cell.text = ""
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        if align_right:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(str(line))
        _set_run_font(run, size_pt=size_pt, bold=(first_bold and idx == 0), color=SSI_INK)


def _split_address(address: str) -> Dict[str, str]:
    raw = re.sub(r"\s+", " ", (address or "").strip())
    if not raw:
        return {"street": "—", "city": "—"}
    match = re.match(r"^(.*?),\s*(\d{4}\s+.+)$", raw)
    if match:
        return {"street": match.group(1).strip() or "—", "city": match.group(2).strip() or "—"}
    match = re.match(r"^(.*?)(\d{4}\s+.+)$", raw)
    if match and match.group(1).strip():
        return {
            "street": match.group(1).replace(",", "").strip() or "—",
            "city": match.group(2).strip() or "—",
        }
    return {"street": raw, "city": "—"}


def _format_date_de(value: Any) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    if re.fullmatch(r"\d{2}\.\d{2}\.\d{4}", raw):
        return raw
    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})", raw)
    if m:
        return f"{m.group(3)}.{m.group(2)}.{m.group(1)}"
    return raw


def _add_heading(doc: Document, text: str, *, level: int = 1) -> None:
    sizes = {1: 16, 2: 13, 3: 11}
    _add_para(
        doc,
        text,
        size_pt=sizes.get(level, 11),
        bold=True,
        color=SSI_NAVY if level <= 2 else SSI_INK,
        space_before=12 if level == 1 else 8,
        space_after=6,
    )


def _add_bullets(doc: Document, items: List[str]) -> None:
    for item in items or []:
        text = str(item or "").strip()
        if not text:
            continue
        try:
            p = doc.add_paragraph(style="List Bullet")
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            run = p.add_run(text)
        except KeyError:
            p = doc.add_paragraph()
            run = p.add_run(f"• {text}")
        _set_run_font(run, size_pt=10, color=SSI_INK)
        p.paragraph_format.space_after = Pt(2)


def _page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _add_image_if_exists(doc: Document, path: Path, *, width_cm: float) -> bool:
    if not path.exists():
        return False
    try:
        doc.add_picture(str(path), width=Cm(width_cm))
        return True
    except Exception:
        return False


def _build_summary_table(
    doc: Document,
    *,
    title: str,
    groups: List[Dict[str, Any]],
    total_label: str,
    show_qty: bool = False,
    show_price: bool = False,
    note: str = "",
) -> None:
    if not groups:
        return
    rows: List[Dict[str, Any]] = []
    if note:
        rows.append({"kind": "note", "text": note})
    pos = 0
    for group in groups:
        group_note = (group.get("note") or "").strip()
        if group_note and group_note != note:
            rows.append({"kind": "note", "text": group_note})
        for item in group.get("items") or []:
            pos += 1
            name = (item.get("name") or "").strip() or f"Position {pos}"
            desc = (item.get("description") or "").strip()
            qty = item.get("qty")
            qty_text = "—"
            if qty is not None and qty != "":
                try:
                    qn_val = float(qty)
                    qty_text = str(int(qn_val)) if qn_val == int(qn_val) else str(qn_val)
                except (TypeError, ValueError):
                    qty_text = str(qty)
            price_text = _money_compact(item.get("amount")) if show_price and item.get("amount") is not None else ""
            rows.append(
                {
                    "kind": "item",
                    "title": f"{pos}. {name}",
                    "desc": desc,
                    "qty": qty_text if show_qty else None,
                    "price": price_text if show_price else None,
                }
            )
        if group.get("total") is not None:
            rows.append(
                {
                    "kind": "total",
                    "label": total_label,
                    "value": _money_compact(group.get("total")),
                }
            )

    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    table.autofit = True
    header = table.rows[0].cells
    header[0].merge(header[1])
    _shade_cell(header[0], "FFED00")
    _set_cell_text(header[0], title, bold=True, size_pt=10)

    for idx, row in enumerate(rows, start=1):
        left, right = table.rows[idx].cells
        kind = row["kind"]
        if kind == "note":
            left.merge(right)
            _set_cell_text(left, row["text"], size_pt=8, color=SSI_MUTED)
        elif kind == "total":
            _set_cell_text(left, row["label"], bold=True, size_pt=10)
            _set_cell_text(right, row["value"], bold=True, size_pt=10, align_right=True)
        else:
            lines = [row["title"]]
            if row.get("desc"):
                lines.append(row["desc"])
            _set_cell_paragraphs(left, lines, first_bold=True, size_pt=9)
            meta_lines: List[str] = []
            if show_price and show_qty:
                meta_lines = [row.get("price") or "—", f'{row.get("qty") or "—"}×']
            elif show_qty:
                meta_lines = [row.get("qty") or "—"]
            elif show_price:
                meta_lines = [row.get("price") or "—"]
            if meta_lines:
                _set_cell_paragraphs(right, meta_lines, first_bold=True, align_right=True, size_pt=9)
            else:
                right.text = ""
    _add_para(doc, "", space_after=8)


def _classify_scope_groups(scope_groups: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
    lic: List[Dict[str, Any]] = []
    lic_opt: List[Dict[str, Any]] = []
    it: List[Dict[str, Any]] = []
    material: List[Dict[str, Any]] = []
    for group in scope_groups or []:
        gid = str(group.get("id") or "").lower()
        title = str(group.get("title") or "").lower()
        if gid == "licenseoptions" or "optionen" in title or title.startswith("optional"):
            lic_opt.append(group)
        elif gid == "material" or "material" in title:
            material.append(group)
        elif gid == "license" or "lizenz" in title or "software" in title:
            lic.append(group)
        elif gid == "it" or "it-" in title or "it " in title or "reise" in title:
            it.append(group)
        else:
            lic.append(group)
    return {"license": lic, "licenseOptions": lic_opt, "it": it, "material": material}


def _append_ssi_signature_controls(doc: Document) -> bool:
    """Hängt die SSI-Unterschriften-Inhaltssteuerelemente (Name/Funktion) aus der Logimat-Vorlage an."""
    from copy import deepcopy

    from docx_export import find_placeholder_template

    try:
        src = Document(str(find_placeholder_template()))
    except Exception:
        return False

    children = list(src.element.body)
    start = None
    end = None
    for i, child in enumerate(children):
        text = _element_text(child).lower()
        if start is None and "freundliche" in text and "gr" in text:
            start = i
        aliases = [(a.get(qn("w:val")) or "") for a in child.iter(qn("w:alias"))]
        if any("Unterschrift" in a or "Funktion" in a for a in aliases):
            end = i
    if end is None:
        return False
    # Nur Gruss + Firma + Unterschriften-Tabelle (nicht den ganzen Vertragsschluss)
    if start is None:
        start = end
        # ein paar Leerabsätze davor mitnehmen
        start = max(0, end - 8)

    # Bevorzugt ab «Freundliche Grüsse»
    greet_idx = None
    for i in range(start, end + 1):
        if "freundliche" in _element_text(children[i]).lower():
            greet_idx = i
            break
    if greet_idx is not None:
        start = greet_idx

    target = doc.element.body
    sect = target.find(qn("w:sectPr"))
    for child in children[start : end + 1]:
        if child.tag == qn("w:sectPr"):
            continue
        cloned = deepcopy(child)
        if sect is not None:
            sect.addprevious(cloned)
        else:
            target.append(cloned)
    return True


def _add_terms(doc: Document, terms: Dict[str, Any], offer_date: str) -> None:
    if not terms or not (terms.get("sections") or []):
        return
    _page_break(doc)
    _add_para(doc, "Rechtliche Bedingungen", size_pt=9, bold=True, color=SSI_MUTED, space_after=2)
    _add_heading(doc, terms.get("title") or "Vertragsbedingungen", level=1)
    meta = f"Version {terms.get('version') or ''} · Angebotsgültigkeit {terms.get('validityDays') or 30} Tage"
    _add_para(doc, meta, size_pt=9, color=SSI_MUTED, space_after=8)
    if terms.get("intro"):
        _add_para(doc, str(terms.get("intro")), size_pt=10, space_after=10)

    for sec in terms.get("sections") or []:
        sid = str(sec.get("id") or "").strip()
        title = str(sec.get("title") or "").strip()
        _add_heading(doc, f"{sid} {title}".strip(), level=2)
        for p in sec.get("paragraphs") or []:
            if p:
                _add_para(doc, str(p), size_pt=10, space_after=6)
        for b in sec.get("bullets") or []:
            if b:
                _add_bullets(doc, [str(b)])
        for p in sec.get("paragraphsAfter") or []:
            if p:
                _add_para(doc, str(p), size_pt=10, space_after=6)
        for sub in sec.get("subsections") or []:
            if not isinstance(sub, dict):
                continue
            _add_heading(doc, str(sub.get("title") or ""), level=3)
            for p in sub.get("paragraphs") or []:
                if p:
                    _add_para(doc, str(p), size_pt=10, space_after=4)
        table = sec.get("table")
        if isinstance(table, dict) and (table.get("headers") or table.get("rows")):
            headers = [str(h) for h in (table.get("headers") or [])]
            rows = table.get("rows") or []
            cols = max(len(headers), max((len(r) for r in rows), default=1), 1)
            t = doc.add_table(rows=1 + len(rows), cols=cols)
            t.style = "Table Grid"
            for i in range(cols):
                _shade_cell(t.rows[0].cells[i], "EEEEEE")
                _set_cell_text(t.rows[0].cells[i], headers[i] if i < len(headers) else "", bold=True, size_pt=8)
            for r_idx, row in enumerate(rows, start=1):
                for c_idx in range(cols):
                    val = str(row[c_idx]) if c_idx < len(row) else ""
                    _set_cell_text(t.rows[r_idx].cells[c_idx], val, size_pt=8)
            _add_para(doc, "", space_after=6)

    closing = terms.get("closing") or {}
    if closing.get("text"):
        _add_para(doc, str(closing.get("text") or ""), size_pt=10, space_before=10, space_after=8)

    # SSI-Unterschriften als echte Word-Inhaltssteuerelemente (Name + Funktion),
    # damit sie in Word erkannt und später nachgepflegt werden können.
    used_sdt = _append_ssi_signature_controls(doc)
    if not used_sdt:
        greet = closing.get("greeting") or "Freundliche Grüsse"
        company = closing.get("company") or "SSI SCHÄFER AG"
        _add_para(doc, f"{greet}\n{company}", size_pt=10, bold=True, space_after=12)
        sigs = closing.get("signatories") or []
        if sigs:
            t = doc.add_table(rows=1, cols=min(2, len(sigs)) or 1)
            t.autofit = True
            for idx, sig in enumerate(sigs[:2]):
                cell = t.rows[0].cells[idx]
                lines = [
                    "______________________________",
                    str(sig.get("name") or ""),
                    str(sig.get("title") or ""),
                ]
                if sig.get("role"):
                    lines.append(str(sig.get("role")))
                if sig.get("email"):
                    lines.append(str(sig.get("email")))
                _set_cell_paragraphs(cell, [ln for ln in lines if ln], first_bold=False, size_pt=9)
            _add_para(doc, "", space_after=10)

    place = f"Schaffhauserstrasse 10, 8213 Neunkirch{(' · ' + offer_date) if offer_date else ''}"
    acc = doc.add_table(rows=1, cols=2)
    _set_cell_paragraphs(acc.rows[0].cells[0], ["Ort / Datum", place], first_bold=True, size_pt=9)
    _set_cell_paragraphs(
        acc.rows[0].cells[1],
        ["Unterschrift / Stempel Kunde", "______________________________"],
        first_bold=True,
        size_pt=9,
    )



def _element_text(el) -> str:
    return " ".join(t.text for t in el.iter(qn("w:t")) if t.text).strip()


def _element_has_page_break(el) -> bool:
    for br in el.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _trim_to_cover_and_philosophy(doc: Document) -> None:
    """Behält Titelseite (Cover) und SSI-Philosophie; Rest der Logimat-Vorlage fällt weg."""
    body = doc.element.body
    children = list(body)
    if not children:
        return

    first_pb = None
    for i, child in enumerate(children):
        if child.tag == qn("w:p") and _element_has_page_break(child):
            first_pb = i
            break
    if first_pb is None:
        first_pb = min(20, len(children) - 1)

    phil_start = None
    for i, child in enumerate(children):
        text = _element_text(child).lower()
        compact = re.sub(r"\s+", "", text)
        if "ihrpartnerssi" in compact or "partnerssischaefer" in compact or "partnerssischäfer" in compact:
            phil_start = i
            break
    if phil_start is None:
        for i, child in enumerate(children):
            text = _element_text(child).strip().lower()
            compact = re.sub(r"\s+", "", text)
            if compact == "philosophie" or compact.startswith("philosophie"):
                phil_start = i
                break
            if "unternehmensprofil" in compact:
                phil_start = i
                break
    # Wenn wir bei PHILOSOPHIE gelandet sind, 1–3 Blöcke davor mitnehmen (Partner/Profil)
    if phil_start is not None:
        for back in range(1, 4):
            idx = phil_start - back
            if idx <= first_pb:
                break
            compact = re.sub(r"\s+", "", _element_text(children[idx]).lower())
            if (
                "partner" in compact
                or "unternehmensprofil" in compact
                or "philosophie" in compact
                or not compact
            ):
                phil_start = idx
            else:
                break

    phil_end = None
    start_at = phil_start if phil_start is not None else first_pb + 1
    for i, child in enumerate(children):
        if i <= start_at:
            continue
        text = _element_text(child).lower()
        if child.tag == qn("w:tbl") and (
            "einleitung" in text
            or "leistungsumfang" in text
            or "optionen" in text
            or "{{" in text
        ):
            phil_end = i
            break
        if "kundenseitige voraussetzungen" in text:
            phil_end = i
            break
    if phil_end is None:
        # Philosophie endet vor dem nächsten Seitenumbruch nach dem Start
        for i, child in enumerate(children):
            if i <= start_at:
                continue
            if _element_has_page_break(child):
                phil_end = i
                break
    if phil_end is None:
        phil_end = len(children)

    keep = set(range(0, first_pb + 1))
    if phil_start is not None:
        keep.update(range(phil_start, phil_end))
    # Sicherstellen, dass nach dem Cover ein Seitenumbruch bleibt
    keep.add(first_pb)

    for i, child in enumerate(children):
        if child.tag == qn("w:sectPr"):
            continue
        if i not in keep:
            body.remove(child)


def _build_ssi_cover_philosophy_document(offer: Dict[str, Any]) -> Document:
    """Logimat-Vorlage: Cover + Philosophie befüllt, restlicher Alt-Inhalt entfernt."""
    from docxtpl import DocxTemplate

    from docx_export import (
        apply_sdt_contact_fields,
        build_template_context,
        find_placeholder_template,
    )

    template_path = find_placeholder_template()
    context = build_template_context(offer)
    tpl = DocxTemplate(str(template_path))
    tpl.render(context, autoescape=True)
    buf = io.BytesIO()
    tpl.save(buf)
    filled = apply_sdt_contact_fields(buf.getvalue(), context)
    doc = Document(io.BytesIO(filled))
    _trim_to_cover_and_philosophy(doc)
    return doc


def build_offer_docx_html(offer: Dict[str, Any]) -> bytes:
    """DOCX: SSI-Cover + Philosophie aus Logimat-Vorlage, danach HTML-Vorschau-Inhalt."""
    if offer.get("kind") != "offer_document":
        raise ValueError("Word-Export ist nur für Gesamtangebote (offer_document) verfügbar.")

    meta = offer.get("meta") or {}
    content = offer.get("content") or {}
    cfg = content.get("configurationSummary") or {}
    summary = offer.get("priceSummary") or {}
    scope_groups = offer.get("scopeGroups") or content.get("scopeGroups") or []
    arch = content.get("architecture") or {}
    req = content.get("requirements") or {}
    terms = content.get("commercialTerms") or {}
    offer_date = meta.get("documentDate") or _format_date_de(meta.get("createdAt"))

    # Basis: SSI-Titelseite + Philosophie aus der Logimat-Vorlage
    try:
        doc = _build_ssi_cover_philosophy_document(offer)
        _page_break(doc)
    except Exception:
        # Fallback ohne Vorlage: bisheriges programmatisches Cover
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        customer = offer.get("customer") or {}
        addr = _split_address(customer.get("address") or "")
        title = str(content.get("title") or "Angebot WAMAS® Lift & Store").replace("Lift & Store", "Lift Store")
        subtitle = content.get("subtitle") or "SSI SCHÄFER · Softwarelösung für Vertical Lift Modules (SSI LOGIMAT®)"
        project_line = meta.get("archiveTitle") or meta.get("projectLabel") or meta.get("offerNumber") or ""
        ssi = offer.get("ssiContacts") or []
        closing_sigs = (terms.get("closing") or {}).get("signatories") or []
        ssi1 = ssi[0] if len(ssi) > 0 else (closing_sigs[0] if closing_sigs else {})
        ssi2 = ssi[1] if len(ssi) > 1 else (closing_sigs[1] if len(closing_sigs) > 1 else {})
        cover_img = ASSETS_DIR / "offer-cover-logimat.jpg"
        if cover_img.exists():
            _add_image_if_exists(doc, cover_img, width_cm=17.4)
        _add_para(doc, title, size_pt=22, bold=True, color=SSI_NAVY, space_before=10, space_after=4)
        _add_para(doc, str(project_line), size_pt=14, bold=True, space_after=2)
        _add_para(doc, offer_date, size_pt=11, color=SSI_MUTED, space_after=4)
        _add_para(doc, str(subtitle), size_pt=10, color=SSI_MUTED, space_after=12)
        cust = doc.add_table(rows=6, cols=2)
        cust.style = "Table Grid"
        for i, (label, value) in enumerate([
            ("Firmenname", customer.get("company") or "—"),
            ("Name", customer.get("contact") or "—"),
            ("Strasse", addr["street"]),
            ("Postleitzahl, Ort", addr["city"]),
            ("Telefon", customer.get("phone") or ""),
            ("E-Mail", customer.get("email") or ""),
        ]):
            _shade_cell(cust.rows[i].cells[0], "F5F5F5")
            _set_cell_text(cust.rows[i].cells[0], label, bold=True, size_pt=9)
            _set_cell_text(cust.rows[i].cells[1], str(value), size_pt=9)
        _page_break(doc)

    # ---- Body (Vorschau-Inhalt) ----
    _add_para(doc, content.get("annexLabel") or "Anhang zur Software", size_pt=9, bold=True, color=SSI_MUTED, space_after=2)
    _add_heading(doc, "Leistungsbeschreibung WAMAS® Lift & Store", level=1)

    _add_heading(doc, "Inhaltsverzeichnis", level=2)
    toc_items = [
        "A1. Umfang WAMAS Lift & Store",
        "A1.1 Standard-Funktionen / Prozesse",
        "A1.2 Gewählte Software-Optionen",
        "1. Leistungsumfang & Preise",
        "A1.3 Bedienoberflächen",
        "A2. Standard-Systemarchitektur",
        "A3. Anforderungen",
        "A4. Zuständigkeiten",
        "A5. Begleitende Dokumente",
        "Vertragsbedingungen (Software)",
    ]
    for item in toc_items:
        _add_para(doc, item, size_pt=10, space_after=2)

    _add_heading(doc, "A1. Umfang WAMAS Lift & Store", level=1)
    if content.get("intro"):
        _add_para(doc, str(content.get("intro")), size_pt=10, space_after=6)
    if content.get("introVariant"):
        _add_para(doc, str(content.get("introVariant")), size_pt=10, bold=True, space_after=8)

    _add_heading(doc, "A1.1 Standard-Funktionen / Prozesse", level=2)
    if content.get("standardLead"):
        _add_para(doc, str(content.get("standardLead")), size_pt=10, space_after=4)
    if content.get("recommendation"):
        _add_para(doc, str(content.get("recommendation")), size_pt=10, space_after=6)
    for idx, fn in enumerate(content.get("standardFunctions") or [], start=1):
        _add_para(doc, f"1.1.{idx} {fn.get('title') or ''}", size_pt=10, bold=True, space_after=2)
        if fn.get("text"):
            _add_para(doc, str(fn.get("text")), size_pt=9, color=SSI_MUTED, space_after=6)

    _add_heading(doc, "A1.2 Gewählte Software-Optionen für WAMAS® Lift & Store", level=2)
    if content.get("machineOptionsLead"):
        _add_para(doc, str(content.get("machineOptionsLead")), size_pt=10, space_after=4)
    _add_para(
        doc,
        "Kurzbeschrieb der gewählten Module. Kalkulierte Positionen und Preise stehen nur unter «Leistungsumfang & Preise».",
        size_pt=9,
        color=SSI_MUTED,
        space_after=6,
    )
    selected_opts = content.get("selectedOptions") or []
    if selected_opts:
        for opt in selected_opts:
            _add_para(doc, str(opt.get("title") or ""), size_pt=10, bold=True, space_after=1)
            if opt.get("text"):
                _add_para(doc, str(opt.get("text")), size_pt=9, color=SSI_MUTED, space_after=6)
    else:
        _add_para(doc, "Keine optionalen Softwaremodule gewählt — Preise siehe Leistungsumfang.", size_pt=9, color=SSI_MUTED)

    _add_heading(doc, "1. Leistungsumfang & Preise", level=1)
    cfg_bits = " · ".join(
        x
        for x in [
            cfg.get("instanceName"),
            f"{cfg.get('instanceCount')}×" if cfg.get("instanceCount") else None,
            f"{cfg.get('deviceCount')} Geräte" if cfg.get("deviceCount") is not None else None,
            f"{cfg.get('zoneCount')} Zonen" if cfg.get("zoneCount") is not None else None,
            f"{cfg.get('openingCount')} Öffnungen" if cfg.get("openingCount") is not None else None,
            "Order Handling" if cfg.get("hasOrderHandling") else "Standalone",
        ]
        if x
    )
    if cfg_bits:
        _add_para(doc, cfg_bits, size_pt=10, bold=True, space_after=4)
    _add_para(
        doc,
        "Je Position eine Zeile; ausgewiesen werden die Bereichstotals und das Gesamttotal (ohne Einzelpreise je Zeile).",
        size_pt=9,
        color=SSI_MUTED,
        space_after=8,
    )

    buckets = _classify_scope_groups(list(scope_groups))
    _build_summary_table(
        doc,
        title="Zusammenfassung – Softwarelizenzen",
        groups=buckets["license"],
        total_label="Total A · Softwarelizenzen",
        show_qty=True,
    )
    _build_summary_table(
        doc,
        title="Zusammenfassung – IT-Aufwand",
        groups=buckets["it"],
        total_label="Total B · IT-Aufwand",
    )
    _build_summary_table(
        doc,
        title="Zusammenfassung – Material",
        groups=buckets["material"],
        total_label="Total C · Material",
    )

    grand_rows = []
    if summary.get("subtotalChf") is not None and summary.get("commercialDiscountChf"):
        grand_rows.append(("Zwischentotal", _money(summary.get("subtotalChf"), "CHF")))
        grand_rows.append(("Projektrabatt", f"− {_money(summary.get('commercialDiscountChf'), 'CHF')}"))
    grand_rows.append(
        (
            "Total netto exkl. MwSt.",
            _money(summary.get("grandTotalChf"), "CHF") if summary.get("grandTotalChf") is not None else "—",
        )
    )
    grand = doc.add_table(rows=1 + len(grand_rows), cols=2)
    grand.style = "Table Grid"
    h = grand.rows[0].cells
    h[0].merge(h[1])
    _shade_cell(h[0], "FFED00")
    _set_cell_text(h[0], "Gesamttotal", bold=True, size_pt=10)
    for i, (label, value) in enumerate(grand_rows, start=1):
        _set_cell_text(grand.rows[i].cells[0], label, bold=True, size_pt=10)
        _set_cell_text(grand.rows[i].cells[1], value, bold=True, size_pt=10, align_right=True)
    _add_para(doc, "", space_after=6)

    if buckets["licenseOptions"]:
        _add_heading(doc, "Zusätzliche Option", level=2)
        _build_summary_table(
            doc,
            title="Optional – Softwarelizenzen",
            groups=buckets["licenseOptions"],
            total_label="Optionen (nicht im Total)",
            show_qty=True,
            show_price=True,
            note="Zusätzliche Option — preislich ausgewiesen, nicht im Gesamttotal enthalten.",
        )

    note = summary.get("note") or "Alle Preise in CHF, exkl. MwSt."
    _add_para(doc, str(note), size_pt=9, color=SSI_MUTED, space_after=10)

    _add_heading(doc, "A1.3 Bedienoberfläche Bediener und Admin Client", level=2)
    clients = content.get("clients") or {}
    _add_para(doc, "Touch Client", size_pt=10, bold=True, space_after=1)
    _add_para(doc, str(clients.get("touch") or ""), size_pt=9, space_after=6)
    _add_para(doc, "Admin Client", size_pt=10, bold=True, space_after=1)
    _add_para(doc, str(clients.get("admin") or ""), size_pt=9, space_after=6)
    if clients.get("showMobile"):
        _add_para(doc, "Mobile Terminal", size_pt=10, bold=True, space_after=1)
        _add_para(doc, str(clients.get("mobile") or ""), size_pt=9, space_after=8)

    _add_heading(doc, f"A2. {arch.get('title') or 'Standard-Systemarchitektur'}", level=1)
    if arch.get("text"):
        _add_para(doc, str(arch.get("text")), size_pt=10, space_after=6)
    arch_path = arch.get("image") or ""
    if isinstance(arch_path, str) and arch_path.startswith("/static/"):
        local = BASE_DIR / arch_path.lstrip("/")
        # strip query
        local = Path(str(local).split("?", 1)[0])
        _add_image_if_exists(doc, local, width_cm=16)
    elif (ASSETS_DIR / "annex" / "architecture.png").exists():
        _add_image_if_exists(doc, ASSETS_DIR / "annex" / "architecture.png", width_cm=16)
    _add_bullets(doc, list(arch.get("legend") or []))

    _add_heading(doc, f"A3. {req.get('title') or 'Anforderungen'}", level=1)
    if req.get("note"):
        _add_para(doc, str(req.get("note")), size_pt=10, space_after=4)
    if content.get("serverProvisioningNote"):
        _add_para(doc, str(content.get("serverProvisioningNote")), size_pt=9, color=SSI_MUTED, space_after=6)
    for label, key in [
        ("Server", "server"),
        ("Desktop / Admin Client", "desktop"),
        ("Touch Client / IPC", "touch"),
    ]:
        items = req.get(key) or []
        if items:
            _add_para(doc, label, size_pt=10, bold=True, space_after=2)
            _add_bullets(doc, list(items))
    if req.get("showMobile") and (req.get("mobile") or []):
        _add_para(doc, "Mobile Terminal", size_pt=10, bold=True, space_after=2)
        _add_bullets(doc, list(req.get("mobile") or []))
    if req.get("networkHighlight"):
        _add_para(doc, f"Netzwerk: {req.get('networkHighlight')}", size_pt=9, bold=True, space_after=8)

    _add_heading(doc, "A4. Zuständigkeiten", level=1)
    _add_heading(doc, "Endabnahme", level=3)
    if content.get("acceptance"):
        _add_para(doc, str(content.get("acceptance")), size_pt=10, space_after=6)
    _add_heading(doc, "Zuständigkeitsmatrix", level=3)
    responsibilities = content.get("responsibilities") or []
    if responsibilities:
        rt = doc.add_table(rows=1 + len(responsibilities), cols=3)
        rt.style = "Table Grid"
        for i, htxt in enumerate(["Aufgabe", "SSI", "Kunde"]):
            _shade_cell(rt.rows[0].cells[i], "EEEEEE")
            _set_cell_text(rt.rows[0].cells[i], htxt, bold=True, size_pt=8)
        for r_idx, row in enumerate(responsibilities, start=1):
            _set_cell_text(rt.rows[r_idx].cells[0], str(row.get("task") or ""), size_pt=8)
            _set_cell_text(rt.rows[r_idx].cells[1], "X" if row.get("ssi") else "", size_pt=8)
            _set_cell_text(rt.rows[r_idx].cells[2], "X" if row.get("customer") else "", size_pt=8)

    _add_heading(doc, "A5. Begleitende Dokumente", level=1)
    if content.get("documentsLead"):
        _add_para(doc, str(content.get("documentsLead")), size_pt=10, space_after=4)
    _add_bullets(doc, list(content.get("documents") or []))
    if content.get("closing"):
        _add_para(doc, str(content.get("closing")), size_pt=10, space_before=6, space_after=8)

    _add_terms(doc, terms, offer_date)

    footer = (
        f"{meta.get('offerNumber') or ''} · {offer_date}"
        f" · gültig bis {_format_date_de(meta.get('validUntil'))}"
    ).strip(" ·")
    _add_para(doc, footer, size_pt=8, color=SSI_MUTED, space_before=16, space_after=0)

    # Falls keine Vertragsbedingungen: Unterschriften-Steuerelemente trotzdem anhängen
    aliases = [
        (a.get(qn("w:val")) or "")
        for a in doc.element.body.iter(qn("w:alias"))
    ]
    if not any("Unterschrift" in a or "Funktion" in a for a in aliases):
        _append_ssi_signature_controls(doc)

    from docx_export import apply_sdt_contact_fields, build_template_context

    out = io.BytesIO()
    doc.save(out)
    # Name-/Funktion-Inhaltssteuerelemente mit Signatur-Kontext befüllen
    return apply_sdt_contact_fields(out.getvalue(), build_template_context(offer))
